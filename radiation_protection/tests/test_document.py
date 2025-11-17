from datetime import timedelta
from pathlib import Path
from typing import TYPE_CHECKING
from unittest import mock

import pytest
from django.utils import timezone
from docx import Document

from certification.certifications.models import QuizResult
from certification.certifications.tests.factories import QuizCertificationFactory
from euphro_auth.tests import factories as auth_factories
from lab.runs.models import Run
from lab.tests import factories as lab_factories
from radiation_protection.certification import get_radioprotection_certification
from radiation_protection.document import (
    _prepare_variables,
    _replace_text_in_paragraph,
    _replace_variables_in_document,
    write_authorization_access_form,
    write_risk_prevention_plan,
)
from radiation_protection.tests import factories as radiation_factories

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject


@pytest.fixture(name="mock_document_paths")
def mock_document_paths_fixture(settings):
    """Create mock document paths for testing."""
    base_path = Path(settings.BASE_DIR) / "radiation_protection" / "assets"
    return {
        "default": base_path / "AGLAE_plan_de_prevention.docx",
        "en": base_path / "AGLAE_plan_de_prevention_english.docx",
    }


@pytest.fixture(name="mock_document")
def mock_document_fixture():
    """Create a mock document with test content."""
    doc = Document()
    doc.add_paragraph("Test document for user_name")
    doc.add_paragraph("Admin: admin_name")
    doc.add_paragraph("Start: run_date_start")
    doc.add_paragraph("End: run_date_end")
    doc.add_paragraph("Certification date: certification_date")
    return doc


@pytest.fixture(name="quiz_result")
def quiz_result_fixture():
    """Create a quiz result for testing."""
    user = auth_factories.StaffUserFactory()
    quiz = QuizCertificationFactory(certification=get_radioprotection_certification())
    return QuizResult.objects.create(
        user=user,
        quiz=quiz,
        is_passed=True,
        score=95,
    )


@pytest.fixture(name="next_user_run")
def next_user_run_fixture():
    """Create a next user run for testing."""
    admin = auth_factories.StaffUserFactory()
    project = lab_factories.ProjectFactory(admin=admin)
    return lab_factories.RunFactory(
        project=project,
        start_date=timezone.now() + timedelta(days=7),
        end_date=timezone.now() + timedelta(days=14),
    )


def test_replace_text_in_paragraph():
    """Test text replacement in a paragraph."""
    doc = Document()
    paragraph = doc.add_paragraph("Hello ${name}!")
    _replace_text_in_paragraph(paragraph, "${name}", "John")
    assert paragraph.text == "Hello John!"


@pytest.mark.django_db
def test_prepare_variables_with_run(quiz_result: QuizResult, next_user_run: Run):
    """Test preparing variables when run is available."""
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user, institution=lab_factories.InstitutionFactory()
    )
    variables = _prepare_variables(participation, next_user_run)

    assert variables["user_full_name"] == quiz_result.user.get_administrative_name()  # type: ignore # pylint: disable=line-too-long
    assert variables["admin_name"] == next_user_run.project.admin.get_administrative_name()  # type: ignore # pylint: disable=line-too-long
    assert variables["run_date_start"] == next_user_run.start_date.strftime(  # type: ignore # pylint: disable=line-too-long
        "%d/%m/%Y"
    )
    assert variables["run_date_end"] == next_user_run.end_date.strftime(  # type: ignore # pylint: disable=line-too-long
        "%d/%m/%Y"
    )
    assert variables["certification_date"] == quiz_result.created.strftime("%d/%m/%Y")


@pytest.mark.django_db
def test_prepare_variables_without_run(quiz_result: QuizResult):
    """Test preparing variables when no run is available."""
    participation = lab_factories.ParticipationFactory(
        user=quiz_result.user, institution=lab_factories.InstitutionFactory()
    )
    variables = _prepare_variables(participation, None)

    assert variables["user_full_name"] == quiz_result.user.get_administrative_name()  # type: ignore  # pylint: disable=line-too-long
    assert variables["admin_name"] == ""
    assert variables["run_date_start"] == ""
    assert variables["run_date_end"] == ""
    assert variables["certification_date"] == quiz_result.created.strftime("%d/%m/%Y")


@pytest.mark.django_db
def test_prepare_variables_without_passed_quiz_result(quiz_result: QuizResult):
    """Test preparing variables when no run is available."""
    with mock.patch(
        "radiation_protection.document.get_user_passed_certification_date"
    ) as mock_date:
        mock_date.return_value = None
        participation = lab_factories.ParticipationFactory(
            user=quiz_result.user, institution=lab_factories.InstitutionFactory()
        )
        variables = _prepare_variables(participation, None)

    assert variables["user_full_name"] == quiz_result.user.get_administrative_name()  # type: ignore # pylint: disable=line-too-long
    assert variables["admin_name"] == ""
    assert variables["run_date_start"] == ""
    assert variables["run_date_end"] == ""
    assert variables["certification_date"] == ""


def test_replace_variables_in_document(mock_document: "DocumentObject"):
    """Test replacing variables in document."""
    variables = {
        "user_name": "John Doe",
        "admin_name": "Admin User",
        "run_date_start": "01/01/2024",
        "run_date_end": "02/01/2024",
        "certification_date": "03/01/2024",
    }

    _replace_variables_in_document(mock_document, variables)

    assert mock_document.paragraphs[0].text == "Test document for John Doe"
    assert mock_document.paragraphs[1].text == "Admin: Admin User"
    assert mock_document.paragraphs[2].text == "Start: 01/01/2024"
    assert mock_document.paragraphs[3].text == "End: 02/01/2024"
    assert mock_document.paragraphs[4].text == "Certification date: 03/01/2024"


@pytest.mark.django_db
def test_prepare_variables_with_employer_and_institution(
    quiz_result: QuizResult, next_user_run: Run
):
    """Test preparing variables with employer and institution."""
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=lab_factories.InstitutionFactory(),
        employer=employer,
    )
    variables = _prepare_variables(participation, next_user_run)

    assert variables["user_full_name"] == quiz_result.user.get_administrative_name()  # type: ignore  # pylint: disable=line-too-long
    assert variables["user_email"] == quiz_result.user.email
    assert variables["admin_name"] == next_user_run.project.admin.get_administrative_name()  # type: ignore  # pylint: disable=line-too-long
    assert variables["run_date_start"] == next_user_run.start_date.strftime("%d/%m/%Y")  # type: ignore  # pylint: disable=line-too-long
    assert variables["run_date_end"] == next_user_run.end_date.strftime("%d/%m/%Y")  # type: ignore  # pylint: disable=line-too-long
    assert variables["certification_date"] == quiz_result.created.strftime("%d/%m/%Y")
    assert variables["institution_name"] == participation.institution.name  # type: ignore # pylint: disable=line-too-long
    assert (
        variables["employer_full_name"]
        == f"{employer.last_name.upper()} {employer.first_name}"
    )
    assert variables["employer_email"] == employer.email
    assert variables["employer_function"] == employer.function


@pytest.mark.django_db
def test_prepare_variables_without_institution(
    quiz_result: QuizResult, next_user_run: Run
):
    """Test preparing variables when institution is None."""
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=None,
        employer=employer,
    )
    variables = _prepare_variables(participation, next_user_run)

    assert variables["user_full_name"] == quiz_result.user.get_administrative_name()  # type: ignore  # pylint: disable=line-too-long
    assert variables["user_email"] == quiz_result.user.email
    assert variables["admin_name"] == next_user_run.project.admin.get_administrative_name()  # type: ignore # pylint: disable=line-too-long
    assert variables["run_date_start"] == next_user_run.start_date.strftime("%d/%m/%Y")  # type: ignore # pylint: disable=line-too-long
    assert variables["run_date_end"] == next_user_run.end_date.strftime("%d/%m/%Y")  # type: ignore # pylint: disable=line-too-long
    assert variables["certification_date"] == quiz_result.created.strftime("%d/%m/%Y")
    assert variables["institution_name"] == ""
    assert (
        variables["employer_full_name"]
        == f"{employer.last_name.upper()} {employer.first_name}"
    )
    assert variables["employer_email"] == employer.email
    assert variables["employer_function"] == employer.function


@pytest.mark.django_db
def test_write_authorization_access_form(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    """Test writing authorization access form."""
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=lab_factories.InstitutionFactory(),
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "authorization_form.docx"
    write_authorization_access_form(risk_prevention_plan, output_path)

    assert output_path.exists()
    # Verify the document can be opened
    doc = Document(str(output_path))
    assert len(doc.paragraphs) > 0


@pytest.mark.django_db
def test_write_risk_prevention_plan(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    """Test writing risk prevention plan."""
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=lab_factories.InstitutionFactory(),
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "risk_prevention_plan.docx"
    write_risk_prevention_plan(risk_prevention_plan, output_path)

    assert output_path.exists()
    # Verify the document can be opened
    doc = Document(str(output_path))
    assert len(doc.paragraphs) > 0


@pytest.mark.django_db
def test_write_authorization_access_form_non_french_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    employer = lab_factories.EmployerFactory()
    institution = lab_factories.InstitutionFactory(country="United States")
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=institution,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "authorization_form_en.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_authorization_access_form(risk_prevention_plan, output_path)

        # Verify _create_document was called with English template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" in str(document_path_arg).lower()


@pytest.mark.django_db
def test_write_authorization_access_form_french_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    employer = lab_factories.EmployerFactory()
    institution = lab_factories.InstitutionFactory(country="France")
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=institution,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "authorization_form_fr.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_authorization_access_form(risk_prevention_plan, output_path)

        # Verify _create_document was called with default template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" not in str(document_path_arg).lower()
        assert "AGLAE_plan_de_prevention.docx" in str(
            document_path_arg
        ) or "Formulaire_Autorisation_Acces" in str(document_path_arg)


@pytest.mark.django_db
def test_write_authorization_access_form_no_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=None,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "authorization_form_default.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_authorization_access_form(risk_prevention_plan, output_path)

        # Verify _create_document was called with default template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" not in str(document_path_arg).lower()


@pytest.mark.django_db
def test_write_risk_prevention_plan_non_french_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    employer = lab_factories.EmployerFactory()
    institution = lab_factories.InstitutionFactory(country="Germany")
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=institution,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "risk_prevention_plan_en.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_risk_prevention_plan(risk_prevention_plan, output_path)

        # Verify _create_document was called with English template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" in str(document_path_arg).lower()


@pytest.mark.django_db
def test_write_risk_prevention_plan_french_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    employer = lab_factories.EmployerFactory()
    institution = lab_factories.InstitutionFactory(country="French")
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=institution,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "risk_prevention_plan_fr.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_risk_prevention_plan(risk_prevention_plan, output_path)

        # Verify _create_document was called with default template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" not in str(document_path_arg).lower()


@pytest.mark.django_db
def test_write_risk_prevention_plan_no_institution(
    quiz_result: QuizResult, next_user_run: Run, tmp_path: Path
):
    """Test writing risk prevention plan with no institution uses default template."""
    employer = lab_factories.EmployerFactory()
    participation = next_user_run.project.participation_set.create(
        user=quiz_result.user,
        institution=None,
        employer=employer,
    )
    risk_prevention_plan = radiation_factories.RiskPreventionPlanFactory(
        participation=participation, run=next_user_run
    )

    output_path = tmp_path / "risk_prevention_plan_default.docx"

    with mock.patch("radiation_protection.document._create_document") as mock_create:
        write_risk_prevention_plan(risk_prevention_plan, output_path)

        # Verify _create_document was called with default template
        mock_create.assert_called_once()
        document_path_arg = mock_create.call_args[1]["document_path"]
        assert "english" not in str(document_path_arg).lower()
