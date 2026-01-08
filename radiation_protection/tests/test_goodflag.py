import tempfile
from pathlib import Path
from unittest import mock

import pytest
import requests
from docx import Document

from radiation_protection.electrical_signature.providers.goodflag import (
    GoodflagAPIError,
    GoodflagClient,
    Step,
    StepType,
    start_workflow,
)


@pytest.fixture(name="goodflag_client")
def goodflag_client_fixture():
    """Create a Goodflag client with mock credentials."""
    return GoodflagClient(
        base_url="https://api.goodflag.com",
        api_token="test_token",
        user_id="test_user_id",
    )


@pytest.fixture(name="sample_steps")
def sample_steps_fixture():
    """Create sample workflow steps."""
    return [
        {
            "step_type": StepType.APPROVAL,
            "recipients": [
                {
                    "email": "user@example.com",
                    "first_name": "John",
                    "last_name": "Doe",
                    "preferred_locale": "fr",
                }
            ],
        },
        {
            "step_type": StepType.SIGNATURE,
            "recipients": [
                {
                    "email": "signer@example.com",
                    "first_name": "Jane",
                    "last_name": "Smith",
                    "preferred_locale": "en",
                }
            ],
        },
    ]


def test_goodflag_client_initialization():
    """Test Goodflag client initialization."""
    client = GoodflagClient(
        base_url="https://api.goodflag.com",
        api_token="test_token",
        user_id="test_user_id",
    )
    assert client.base_url == "https://api.goodflag.com/api"
    assert client.api_token == "test_token"
    assert client.user_id == "test_user_id"
    assert "Bearer test_token" in client.headers["Authorization"]


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.app_settings",
    mock.MagicMock(
        GOODFLAG_API_BASE=None, GOODFLAG_API_TOKEN=None, GOODFLAG_USER_ID=None
    ),
)
def test_goodflag_client_missing_credentials():
    """Test Goodflag client raises error when credentials are missing."""
    with pytest.raises(ValueError, match="GOODFLAG_API_BASE"):
        GoodflagClient(base_url=None, api_token="token", user_id="user")

    with pytest.raises(ValueError, match="GOODFLAG_API_TOKEN"):
        GoodflagClient(base_url="http://api.com", api_token=None, user_id="user")

    with pytest.raises(ValueError, match="GOODFLAG_USER_ID"):
        GoodflagClient(base_url="http://api.com", api_token="token", user_id=None)


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_create_workflow(mock_request, goodflag_client: GoodflagClient, sample_steps):
    """Test creating a workflow."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {"id": "workflow_123"}
    mock_response.ok = True
    mock_request.return_value = mock_response

    result = goodflag_client.create_workflow(
        name="Test Workflow",
        steps=sample_steps,
    )

    assert result["id"] == "workflow_123"
    mock_request.assert_called_once()
    call_kwargs = mock_request.call_args[1]
    assert call_kwargs["json"]["name"] == "Test Workflow"
    assert len(call_kwargs["json"]["steps"]) == 2

    # Verify preferredLocale is included for recipients
    step_0_recipient = call_kwargs["json"]["steps"][0]["recipients"][0]
    assert step_0_recipient["preferredLocale"] == "fr"

    step_1_recipient = call_kwargs["json"]["steps"][1]["recipients"][0]
    assert step_1_recipient["preferredLocale"] == "en"


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_start_workflow_method(mock_request, goodflag_client: GoodflagClient):
    """Test starting a workflow."""
    mock_response = mock.Mock()
    mock_response.ok = True
    mock_request.return_value = mock_response

    goodflag_client.start_workflow("workflow_123")

    mock_request.assert_called_once()
    call_kwargs = mock_request.call_args[1]
    assert call_kwargs["json"]["workflowStatus"] == "started"


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_upload_document(mock_request, goodflag_client: GoodflagClient):
    """Test uploading a document."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {
        "documents": [{"id": "doc_123"}],
    }
    mock_response.ok = True
    mock_request.return_value = mock_response

    # Create a temporary document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc = Document()
        doc.add_paragraph("Test document")
        doc.save(temp_file.name)

        try:
            result = goodflag_client.upload_document("workflow_123", temp_file.name)
            assert result["documents"][0]["id"] == "doc_123"
            mock_request.assert_called_once()
        finally:
            Path(temp_file.name).unlink()


def test_upload_document_file_not_found(goodflag_client: GoodflagClient):
    """Test upload document with non-existent file."""
    with pytest.raises(GoodflagAPIError, match="does not exist"):
        goodflag_client.upload_document("workflow_123", "/nonexistent/file.pdf")


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_create_viewer(mock_request, goodflag_client: GoodflagClient):
    """Test creating a viewer."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {"viewer_id": "viewer_123"}
    mock_response.ok = True
    mock_request.return_value = mock_response

    result = goodflag_client.create_viewer("doc_123")

    assert result["viewer_id"] == "viewer_123"
    mock_request.assert_called_once()
    call_kwargs = mock_request.call_args[1]
    assert "redirectUrl" in call_kwargs["json"]
    assert "expired" in call_kwargs["json"]


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_get_workflow_status(mock_request, goodflag_client: GoodflagClient):
    """Test getting workflow status."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {"status": "completed"}
    mock_response.ok = True
    mock_request.return_value = mock_response

    result = goodflag_client.get_workflow_status("workflow_123")

    assert result["status"] == "completed"
    mock_request.assert_called_once()


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_download_document(
    mock_request, goodflag_client: GoodflagClient, tmp_path: Path
):
    """Test downloading a document."""
    mock_response = mock.Mock()
    mock_response.content = b"document content"
    mock_response.ok = True
    mock_request.return_value = mock_response

    output_path = tmp_path / "downloaded.pdf"
    goodflag_client.download_document("doc_123", str(output_path))

    assert output_path.exists()
    assert output_path.read_bytes() == b"document content"


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_api_error_handling(mock_request, goodflag_client: GoodflagClient):
    """Test API error handling."""
    mock_response = mock.Mock()
    mock_response.status_code = 400
    mock_response.json.return_value = {"error": "Bad request"}
    mock_request.side_effect = requests.exceptions.HTTPError(response=mock_response)

    with pytest.raises(GoodflagAPIError) as exc_info:
        goodflag_client.create_workflow("Test", [])

    assert exc_info.value.status_code == 400
    assert exc_info.value.response_data == {"error": "Bad request"}


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.GoodflagClient"
)
def test_start_workflow_function(mock_client_class, sample_steps):
    """Test the start_workflow function."""
    mock_client = mock.Mock()
    mock_client_class.return_value = mock_client

    mock_client.create_workflow.return_value = {"id": "workflow_123"}
    mock_client.upload_document.return_value = {
        "documents": [{"id": "doc_123"}, {"id": "doc_456"}]
    }
    mock_client.create_viewer.return_value = {"viewer_id": "viewer_123"}

    # Create a temporary document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc = Document()
        doc.add_paragraph("Test document")
        doc.save(temp_file.name)

        try:
            workflow_id = start_workflow(
                workflow_name="Test Workflow",
                steps=sample_steps,
                document_path=temp_file.name,
            )

            assert workflow_id == "workflow_123"
            mock_client.create_workflow.assert_called_once()
            mock_client.upload_document.assert_called_once()
            mock_client.start_workflow.assert_called_once_with("workflow_123")
            # Verify create_viewer was called for each document
            assert mock_client.create_viewer.call_count == 2
        finally:
            Path(temp_file.name).unlink()


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_create_workflow_with_default_preferred_locale(
    mock_request, goodflag_client: GoodflagClient
):
    """Test that preferredLocale defaults to 'fr' when not provided."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {"id": "workflow_123"}
    mock_response.ok = True
    mock_request.return_value = mock_response

    # Create steps without preferred_locale
    steps: list[Step] = [
        {
            "step_type": StepType.APPROVAL,
            "recipients": [
                {
                    "email": "user@example.com",
                    "first_name": "John",
                    "last_name": "Doe",
                }
            ],
        },
    ]

    goodflag_client.create_workflow(name="Test Workflow", steps=steps)

    call_kwargs = mock_request.call_args[1]
    recipient = call_kwargs["json"]["steps"][0]["recipients"][0]
    # Should default to 'fr'
    assert recipient["preferredLocale"] == "fr"


@mock.patch(
    "radiation_protection.electrical_signature.providers.goodflag.requests.request"
)
def test_create_workflow_with_multiple_recipients_different_locales(
    mock_request, goodflag_client: GoodflagClient
):
    """Test workflow creation with multiple recipients
    having different preferred locales."""
    mock_response = mock.Mock()
    mock_response.json.return_value = {"id": "workflow_123"}
    mock_response.ok = True
    mock_request.return_value = mock_response

    steps: list[Step] = [
        {
            "step_type": StepType.SIGNATURE,
            "recipients": [
                {
                    "email": "user1@example.com",
                    "first_name": "John",
                    "last_name": "Doe",
                    "preferred_locale": "en",
                },
                {
                    "email": "user2@example.com",
                    "first_name": "Marie",
                    "last_name": "Dupont",
                    "preferred_locale": "fr",
                },
            ],
        },
    ]

    goodflag_client.create_workflow(name="Test Workflow", steps=steps)

    call_kwargs = mock_request.call_args[1]
    recipients = call_kwargs["json"]["steps"][0]["recipients"]
    assert recipients[0]["preferredLocale"] == "en"
    assert recipients[1]["preferredLocale"] == "fr"
