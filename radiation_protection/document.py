import logging
from pathlib import Path
from typing import TYPE_CHECKING

from django.conf import settings
from docx import Document
from docx.text.paragraph import Paragraph

from lab.participations.models import Participation
from lab.runs.models import Run

from .certification import get_user_passed_certification_date
from .models import RiskPreventionPlan

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject

BASE_PATH = Path(settings.BASE_DIR) / "radiation_protection" / "assets"

logger = logging.getLogger(__name__)


def _get_risk_prevention_document_paths() -> dict[str, Path]:
    """Get the path to the radiation protection template document."""
    return {
        settings.DEFAULT_LOCALE: BASE_PATH / "AGLAE_plan_de_prevention.docx",
        "en": BASE_PATH / "AGLAE_plan_de_prevention_english.docx",
    }


def _get_authorization_access_form_path() -> dict[str, Path]:
    """Get the path to the radiation protection authorization access form."""
    return {
        settings.DEFAULT_LOCALE: BASE_PATH
        / "Formulaire_Autorisation_Acces_zone surveillee_ext_AGLAE.docx",
        "en": BASE_PATH
        / "Formulaire_Autorisation_Acces_zone surveillee_ext_AGLAE_english.docx",
    }


def _replace_text_in_paragraph(paragraph: Paragraph, key: str, value: str) -> None:
    """Replace text in a paragraph, even if the key is split across runs,
    preserving formatting.
    """
    # Join all run texts to find the full key
    full_text = "".join(run.text for run in paragraph.runs)
    if key not in full_text:
        return

    # Replace the key with the value in the full text
    new_text = full_text.replace(key, value)

    # Remove all runs except the first
    for _ in range(len(paragraph.runs) - 1):
        paragraph.runs[-1].clear()  # Remove text and formatting
        paragraph._element.remove(  # pylint: disable=protected-access
            paragraph.runs[-1]._element  # pylint: disable=protected-access
        )

    # Set the text of the first run to the new text
    paragraph.runs[0].text = new_text


def _prepare_variables(
    participation: Participation, next_user_run: Run | None
) -> dict[str, str]:
    """Prepare variables for document template replacement (base)."""
    user = participation.user
    passed_certification_date = get_user_passed_certification_date(user)
    if not passed_certification_date:
        logger.warning(
            "User %s does not have a passed radiation protection certification. %s",
            user.email,
            f"Run: {next_user_run.label}" if next_user_run else "",
        )

    variables: dict[str, str] = {
        "certification_date": (
            passed_certification_date.strftime("%d/%m/%Y")
            if passed_certification_date
            else ""
        ),
        "admin_name": "",
        "user_full_name": participation.user.get_administrative_name(),
        "user_email": participation.user.email,
        "run_date_start": "",
        "run_date_end": "",
        "institution_name": "",
        "employer_full_name": "",
        "employer_email": "",
        "employer_function": "",
    }

    if next_user_run:
        admin = next_user_run.project.admin
        if admin:
            variables["admin_name"] = admin.get_administrative_name()

        if next_user_run.start_date:
            variables["run_date_start"] = next_user_run.start_date.strftime("%d/%m/%Y")

        if next_user_run.end_date:
            variables["run_date_end"] = next_user_run.end_date.strftime("%d/%m/%Y")

    if participation.institution:
        variables["institution_name"] = participation.institution.name

    employer = participation.employer

    if employer:
        variables["employer_full_name"] = (
            f"{employer.last_name.upper()} {employer.first_name}"
        )
        variables["employer_email"] = employer.email
        variables["employer_function"] = employer.function

    return variables


def _replace_variables_in_document(
    doc: "DocumentObject", variables: dict[str, str]
) -> None:
    """Replace variables in document paragraphs and tables."""
    # Replace variables in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            _replace_text_in_paragraph(paragraph, key, value)

    # Replace variables in tables
    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        _replace_text_in_paragraph(paragraph, key, value)


def _create_document(
    document_path: Path,
    write_path: Path,
    participation: Participation,
    next_user_run: Run | None,
):
    # Create a copy of the template
    doc = Document(str(document_path))

    # Prepare variables for replacement
    variables = _prepare_variables(participation, next_user_run)

    # Replace variables in document
    _replace_variables_in_document(doc, variables)

    # Write the document to the specified path
    doc.save(str(write_path))


def write_authorization_access_form(plan: RiskPreventionPlan, write_path: Path):
    document_dict = _get_authorization_access_form_path()
    institution = plan.participation.institution
    if (
        institution
        and institution.country
        and institution.country.lower() not in ["france", "french"]
    ):
        document_path = document_dict.get("en", document_dict[settings.DEFAULT_LOCALE])
    else:
        document_path = document_dict[settings.DEFAULT_LOCALE]
    _create_document(
        document_path=document_path,
        write_path=write_path,
        participation=plan.participation,
        next_user_run=plan.run,
    )


def write_risk_prevention_plan(plan: RiskPreventionPlan, write_path: Path):
    document_dict = _get_risk_prevention_document_paths()
    institution = plan.participation.institution
    if institution and institution.get_administrative_locale() == "en":
        document_path = document_dict.get("en", document_dict[settings.DEFAULT_LOCALE])
    else:
        document_path = document_dict[settings.DEFAULT_LOCALE]
    _create_document(
        document_path=document_path,
        write_path=write_path,
        participation=plan.participation,
        next_user_run=plan.run,
    )
